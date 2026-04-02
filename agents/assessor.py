"""
Agent 2: Resume Assessor
Acts as a senior recruiter reviewing your resume against a job description.
Identifies missing keywords, weak bullets, ATS issues, and scores the match.
"""
import json
import re
import time
from pathlib import Path

import ollama
from docx import Document

import config
from config import setup_logging
from db import DatabaseManager

logger = setup_logging("jobbot.assessor")


SYSTEM_PROMPT = """You are a senior technical recruiter who reviews 200 resumes per day for \
software engineering and data/ML/AI roles. You are brutally honest, efficient, and focused on \
ATS optimization. Your assessments directly affect hiring decisions. You know exactly what \
hiring managers look for and what causes resumes to be filtered out before a human even reads them."""

USER_PROMPT_TEMPLATE = """Review this resume against the job description below. Provide a structured \
gap analysis.

RESUME:
{resume_text}

JOB DESCRIPTION:
{job_description}

REQUIRED ATS KEYWORDS FROM JD:
{keywords_json}

Respond ONLY with a valid JSON object following this exact schema (no markdown, no explanation, no <think> tags):
{{
  "match_score": <integer 0-100>,
  "missing_keywords": [<keywords from JD not found in resume>],
  "weak_bullets": [
    {{"original": "<bullet text>", "reason": "<why it's weak>"}}
  ],
  "missing_sections": [<sections the JD expects but resume lacks, e.g. "Certifications">],
  "ats_issues": [<formatting or structural issues that hurt ATS parsing>],
  "strengths": [<top 3 things the resume does well for this role>],
  "critical_gaps": [<top 3 most important things to fix, ordered by priority>],
  "recommended_title": "<suggested job title to use on resume for this application>"
}}"""


def _read_docx_text(docx_path: Path) -> str:
    """Extract all text from a DOCX file including tables."""
    doc = Document(str(docx_path))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts)


def _call_llm_with_retry(system: str, user: str, max_tokens: int = 2000) -> tuple[dict, dict]:
    """Call Ollama and return (parsed_json, usage_dict). Retries on bad JSON."""
    original_user = user  # Preserve original prompt to avoid cumulative mutation
    for attempt in range(config.MAX_RETRIES):
        try:
            # On retry, prepend the retry note fresh to the original prompt
            if attempt > 0:
                current_user = (
                    "IMPORTANT: Your last response was not valid JSON. "
                    "Return ONLY the JSON object, no markdown, no <think> blocks.\n\n"
                    + original_user
                )
            else:
                current_user = original_user

            resp = ollama.chat(
                model=config.OLLAMA_MODEL,
                messages=[
                    {"role": "system", "content": system},
                    {"role": "user", "content": current_user},
                ],
                options={"temperature": 0.2, "num_predict": max_tokens},
            )
            raw = resp["message"]["content"].strip()
            # Strip <think>...</think> blocks emitted by qwen3 in reasoning mode
            raw = re.sub(r"<think>.*?</think>", "", raw, flags=re.DOTALL).strip()
            if raw.startswith("```"):
                raw = raw.split("```")[1].lstrip("json").strip()
            usage = {
                "input_tokens": resp.get("prompt_eval_count", 0),
                "output_tokens": resp.get("eval_count", 0),
            }
            return json.loads(raw), usage
        except json.JSONDecodeError as e:
            logger.warning("LLM returned invalid JSON (attempt %d/%d): %s", attempt + 1, config.MAX_RETRIES, e)
            if attempt == config.MAX_RETRIES - 1:
                raise ValueError(f"LLM returned non-JSON after {config.MAX_RETRIES} attempts")
            time.sleep(2)
        except Exception as e:
            logger.error("LLM call error (attempt %d/%d): %s", attempt + 1, config.MAX_RETRIES, e)
            if attempt == config.MAX_RETRIES - 1:
                raise
            time.sleep(3)
    raise RuntimeError("LLM call failed after all retries")


class ResumeAssessor:
    def __init__(self, db: DatabaseManager):
        self.db = db

    def assess(self, job_id: int) -> dict:
        """Run gap analysis for a job. Returns the gap_analysis dict."""
        job = self.db.get_job(job_id)
        if not job:
            raise ValueError(f"Job {job_id} not found")

        resume_path = config.MASTER_RESUME_PATH
        if not resume_path.exists():
            raise FileNotFoundError(
                f"Master resume not found at {resume_path}. "
                "Place your resume at resumes/master_resume.docx"
            )

        logger.info("Assessing job %d: %s @ %s", job_id, job.get("job_title"), job.get("company"))
        resume_text = _read_docx_text(resume_path)
        if not resume_text.strip():
            raise ValueError("Could not extract text from master_resume.docx")

        if len(resume_text) > 8000:
            logger.warning("Resume text truncated from %d to 8000 chars for job %d", len(resume_text), job_id)
        resume_text = resume_text[:8000]

        job_description = job.get("job_description") or ""
        keywords_raw = job.get("keywords") or "[]"
        if isinstance(keywords_raw, str):
            try:
                keywords = json.loads(keywords_raw)
            except (json.JSONDecodeError, TypeError) as e:
                logger.warning("Failed to parse keywords JSON for job %d, defaulting to empty list: %s", job_id, e)
                keywords = []
        else:
            keywords = keywords_raw

        user_prompt = USER_PROMPT_TEMPLATE.format(
            resume_text=resume_text,
            job_description=job_description[:5000],
            keywords_json=json.dumps(keywords),
        )

        gap_analysis, usage = _call_llm_with_retry(SYSTEM_PROMPT, user_prompt, max_tokens=2000)
        match_score = int(gap_analysis.get("match_score", 0))

        self.db.update_job_gap_analysis(job_id, gap_analysis, match_score)
        self.db.log_action(job_id, "assess", "completed", {"match_score": match_score, "tokens": usage})

        logger.info("Assessment complete — job %d match score: %d%%", job_id, match_score)
        return gap_analysis

    def assess_batch(self, job_ids: list[int]) -> dict[int, dict]:
        """Assess multiple jobs sequentially."""
        results = {}
        for job_id in job_ids:
            try:
                results[job_id] = self.assess(job_id)
            except Exception as e:
                logger.error("Assessment failed for job %d: %s", job_id, e)
                results[job_id] = {"error": str(e)}
        return results
