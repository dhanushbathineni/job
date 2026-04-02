"""
Agent 3: Resume Composer V2 (Iterative with LangGraph)
Uses a multi-step workflow where composer and assessor collaborate to iteratively
improve the resume until it meets quality thresholds.

Workflow:
1. Parse resume and job description
2. Generate initial draft using CAR/STAR methods
3. Assess the draft against the job (gap analysis)
4. If score < threshold or critical gaps remain, revise
5. Validate ATS compatibility
6. Output final DOCX
"""
import json
import re
from pathlib import Path
from typing import TypedDict, Literal, Annotated

from langgraph.graph import StateGraph, END
import ollama
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import config
from config import setup_logging
from db import DatabaseManager
from agents.assessor import _read_docx_text, _call_llm_with_retry

logger = setup_logging("jobbot.composer_v2")

# =============================================================================
# Configuration
# =============================================================================

MATCH_SCORE_THRESHOLD = 75  # Stop iterating if score >= this
MAX_REVISION_ROUNDS = 3     # Maximum revision iterations
ATS_CHECK_ENABLED = True    # Enable ATS validation step


# =============================================================================
# Prompts - Based on proven resume writing frameworks
# =============================================================================

INITIAL_DRAFT_SYSTEM = """You are a professional resume writer with 15+ years of experience in talent acquisition.
You write compelling, results-driven bullet points using the CAR method (Challenge-Action-Result).
You NEVER fabricate experience - you reframe and emphasize existing achievements.

ATS Writing Rules:
- Use standard section headers: "Experience", "Education", "Skills", "Projects"
- No tables, text boxes, graphics, or columns
- Use simple bullet points with standard characters
- Include keywords naturally, not stuffed
- Use reverse chronological order
- Start bullets with strong action verbs (Led, Architected, Built, Optimized, Reduced)
"""

INITIAL_DRAFT_USER = """Rewrite this resume for the target job using the CAR method.

MASTER RESUME:
{resume_text}

TARGET JOB: {job_title} at {company}

JOB DESCRIPTION:
{job_description}

TOP KEYWORDS FROM JD (must include naturally):
{keywords}

GAP ANALYSIS FROM ASSESSOR:
{gap_analysis}

Requirements:
1. Professional Summary: 2-3 sentences, results-driven, mentions target role
2. Experience Bullets: Rewrite using CAR method - Challenge, Action, Result
3. Include quantifiable achievements (%, $, time saved, people led)
4. Incorporate missing keywords naturally
5. Use the recommended_title from gap analysis as the headline

Respond ONLY with valid JSON matching this schema:
{{
  "name": str,
  "contact": {{"email": str, "phone": str, "linkedin": str, "github": str, "location": str}},
  "headline": str,
  "summary": str,
  "skills": {{"languages": [str], "frameworks": [str], "tools": [str], "platforms": [str]}},
  "experience": [{{"title": str, "company": str, "location": str, "start_date": str, "end_date": str, "bullets": [str]}}],
  "education": [{{"degree": str, "institution": str, "year": str, "details": str}}],
  "projects": [{{"name": str, "description": str, "technologies": [str], "link": str}}],
  "certifications": [str]
}}
"""

REVISION_SYSTEM = """You are a resume coach improving a draft based on assessment feedback.
Focus on:
1. Strengthening weak bullet points with quantifiable results
2. Adding missing keywords from the job description
3. Improving the professional summary to be more specific and impactful
4. Fixing any ATS formatting issues identified

For each bullet point, ask: "So what?" - what was the measurable impact?
Transform vague duties into accomplishments.

Example transformations:
- "Responsible for backend development" -> "Architected microservices backend serving 50K+ daily users, reducing API latency by 40%"
- "Worked on machine learning models" -> "Deployed fraud detection model reducing false positives by 28%, saving $200K annually"
"""

REVISION_USER = """Improve this resume draft based on the assessment feedback.

CURRENT DRAFT:
{current_draft}

ASSESSMENT FEEDBACK:
{assessment}

MISSING KEYWORDS (must incorporate):
{missing_keywords}

WEAK BULLETS TO FIX:
{weak_bullets}

ATS ISSUES TO ADDRESS:
{ats_issues}

Provide the revised resume as valid JSON with the same schema as the initial draft.
Focus on making each bullet point more impactful and quantifiable.
"""

ATS_CHECK_SYSTEM = """You are an ATS (Applicant Tracking System) validation expert.
Review this resume for formatting and content issues that cause parsing failures.

Common ATS failures:
- Tables used for layout
- Text boxes or floating elements
- Headers/footers containing contact info
- Graphics, icons, or images
- Multi-column layouts
- Unusual section headers (use: Summary, Experience, Education, Skills)
- Special characters besides standard bullets (•)
- Inconsistent date formats
- Missing standard sections

Respond with JSON:
{{
  "ats_safe": bool,
  "issues": [{{"type": str, "description": str, "severity": "critical"|"warning"}}],
  "recommendations": [str]
}}
"""

ATS_CHECK_USER = """Review this resume for ATS compatibility.

RESUME CONTENT:
{resume_json}

Check for:
1. Formatting that breaks ATS parsers
2. Missing standard sections
3. Keyword density (should be 2-4% for key skills)
4. Section header clarity
5. Date format consistency
6. Contact information placement

Respond with valid JSON identifying any issues.
"""


# =============================================================================
# LangGraph State Definition
# =============================================================================

class ComposerState(TypedDict):
    """State passed through the workflow graph."""
    job_id: int
    job_title: str
    company: str
    job_description: str
    keywords: list[str]
    resume_text: str
    gap_analysis: dict
    current_draft: dict
    assessment: dict
    revision_count: int
    ats_report: dict
    final_score: int
    stop_reason: str


# =============================================================================
# Workflow Nodes
# =============================================================================

def generate_initial_draft(state: ComposerState) -> ComposerState:
    """Generate the first draft of the tailored resume."""
    logger.info("Generating initial draft for job: %s @ %s", state["job_title"], state["company"])

    user_prompt = INITIAL_DRAFT_USER.format(
        resume_text=state["resume_text"][:8000],
        job_title=state["job_title"],
        company=state["company"],
        job_description=state["job_description"][:5000],
        keywords=", ".join(state["keywords"][:20]),
        gap_analysis=json.dumps(state["gap_analysis"], indent=2),
    )

    draft, usage = _call_llm_with_retry(
        system=INITIAL_DRAFT_SYSTEM,
        user=user_prompt,
        max_tokens=4000,
    )

    logger.info("Initial draft generated (input: %d tokens, output: %d tokens)",
                usage.get("input_tokens", 0), usage.get("output_tokens", 0))

    state["current_draft"] = draft
    return state


def assess_draft(state: ComposerState) -> ComposerState:
    """Assess the current draft against the job description."""
    logger.info("Assessing draft quality...")

    # Build a mini assessor prompt
    assess_user = f"""Review this resume draft against the job description.

RESUME DRAFT:
{json.dumps(state["current_draft"], indent=2)}

JOB DESCRIPTION:
{state["job_description"][:4000]}

REQUIRED KEYWORDS:
{", ".join(state["keywords"][:20])}

Respond ONLY with valid JSON:
{{
  "match_score": int (0-100),
  "missing_keywords": [str],
  "weak_bullets": [{{"original": str, "reason": str}}],
  "strengths": [str],
  "critical_gaps": [str],
  "ats_issues": [str],
  "recommendations": [str]
}}
"""

    assessment, usage = _call_llm_with_retry(
        system="You are a senior technical recruiter. Be honest and specific about gaps.",
        user=assess_user,
        max_tokens=2500,
    )

    state["assessment"] = assessment
    state["final_score"] = assessment.get("match_score", 0)

    logger.info("Draft assessment: score=%d, missing_keywords=%d",
                state["final_score"], len(assessment.get("missing_keywords", [])))

    return state


def should_revise(state: ComposerState) -> Literal["revise", "finalize"]:
    """Decision node: continue revising or move to finalization."""
    score = state["final_score"]
    rounds = state["revision_count"]

    if score >= MATCH_SCORE_THRESHOLD:
        state["stop_reason"] = f"Score threshold met ({score} >= {MATCH_SCORE_THRESHOLD})"
        return "finalize"

    if rounds >= MAX_REVISION_ROUNDS:
        state["stop_reason"] = f"Max revision rounds reached ({rounds})"
        return "finalize"

    state["revision_count"] = rounds + 1
    return "revise"


def revise_draft(state: ComposerState) -> ComposerState:
    """Revise the draft based on assessment feedback."""
    logger.info("Revising draft (round %d/%d)...", state["revision_count"], MAX_REVISION_ROUNDS)

    assessment = state["assessment"]

    user_prompt = REVISION_USER.format(
        current_draft=json.dumps(state["current_draft"], indent=2),
        assessment=json.dumps(assessment, indent=2),
        missing_keywords=", ".join(assessment.get("missing_keywords", [])[:15]),
        weak_bullets=json.dumps(assessment.get("weak_bullets", [])[:5]),
        ats_issues=", ".join(assessment.get("ats_issues", [])[:5]),
    )

    revised, usage = _call_llm_with_retry(
        system=REVISION_SYSTEM,
        user=user_prompt,
        max_tokens=4000,
    )

    state["current_draft"] = revised
    logger.info("Revision complete (input: %d tokens, output: %d tokens)",
                usage.get("input_tokens", 0), usage.get("output_tokens", 0))

    return state


def check_ats_compatibility(state: ComposerState) -> ComposerState:
    """Validate ATS compatibility of the final draft."""
    if not ATS_CHECK_ENABLED:
        state["ats_report"] = {"ats_safe": True, "issues": [], "recommendations": []}
        return state

    logger.info("Running ATS compatibility check...")

    user_prompt = ATS_CHECK_USER.format(
        resume_json=json.dumps(state["current_draft"], indent=2),
    )

    ats_report, usage = _call_llm_with_retry(
        system=ATS_CHECK_SYSTEM,
        user=user_prompt,
        max_tokens=2000,
    )

    state["ats_report"] = ats_report

    critical_issues = [i for i in ats_report.get("issues", []) if i.get("severity") == "critical"]
    if critical_issues:
        logger.warning("ATS check found %d critical issues", len(critical_issues))
    else:
        logger.info("ATS check passed - no critical issues")

    return state


def finalize(state: ComposerState) -> ComposerState:
    """Final preparation before DOCX generation."""
    logger.info("Finalizing resume - %s", state["stop_reason"])
    return state


# =============================================================================
# Build the Workflow Graph
# =============================================================================

def build_workflow() -> StateGraph:
    """Construct the LangGraph workflow."""
    workflow = StateGraph(ComposerState)

    # Add nodes
    workflow.add_node("generate_draft", generate_initial_draft)
    workflow.add_node("assess", assess_draft)
    workflow.add_node("revise", revise_draft)
    workflow.add_node("check_ats", check_ats_compatibility)
    workflow.add_node("finalize", finalize)

    # Define edges
    workflow.set_entry_point("generate_draft")
    workflow.add_edge("generate_draft", "assess")

    # Conditional edge after assessment
    workflow.add_conditional_edges(
        "assess",
        should_revise,
        {
            "revise": "revise",
            "finalize": "check_ats",
        },
    )

    # Revision loops back to assessment
    workflow.add_edge("revise", "assess")

    # Final steps
    workflow.add_edge("check_ats", "finalize")
    workflow.add_edge("finalize", END)

    return workflow.compile()


# =============================================================================
# DOCX Builder (same ATS-safe rules as original)
# =============================================================================

def _add_horizontal_rule(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2C3E7A")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_font(run, size: float, bold: bool = False, color: str | None = None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)


def _add_section_header(doc: Document, text: str):
    para = doc.add_paragraph()
    run = para.add_run(text.upper())
    _set_font(run, 10.5, bold=True, color="2C3E7A")
    _add_horizontal_rule(para)
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(2)
    return para


def _add_bullet(doc: Document, text: str):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.15)
    para.paragraph_format.space_after = Pt(1)
    run = para.add_run(f"\u2022  {text}")
    _set_font(run, 9.5)
    return para


def _build_docx(data: dict, output_path: Path) -> Path:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    doc.styles["Normal"].paragraph_format.space_before = Pt(0)
    doc.styles["Normal"].paragraph_format.space_after = Pt(0)

    # Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(data.get("name", ""))
    _set_font(name_run, 16, bold=True, color="1A1A2E")
    name_para.paragraph_format.space_after = Pt(2)

    # Headline
    if data.get("headline"):
        hl_para = doc.add_paragraph()
        hl_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hl_run = hl_para.add_run(data["headline"])
        _set_font(hl_run, 10.5, color="2C3E7A")
        hl_para.paragraph_format.space_after = Pt(3)

    # Contact line
    contact = data.get("contact", {})
    contact_parts = []
    for key in ("email", "phone", "location"):
        if contact.get(key):
            contact_parts.append(contact[key])
    if contact.get("linkedin"):
        contact_parts.append(f"LinkedIn: {contact['linkedin']}")
    if contact.get("github"):
        contact_parts.append(f"GitHub: {contact['github']}")
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_para.add_run("  |  ".join(contact_parts))
    _set_font(contact_run, 9, color="444444")
    contact_para.paragraph_format.space_after = Pt(4)

    # Summary
    if data.get("summary"):
        _add_section_header(doc, "Professional Summary")
        p = doc.add_paragraph()
        r = p.add_run(data["summary"])
        _set_font(r, 9.5)
        p.paragraph_format.space_after = Pt(3)

    # Skills
    skills = data.get("skills", {})
    if any(skills.values()):
        _add_section_header(doc, "Technical Skills")
        skill_lines = [
            ("Languages", skills.get("languages", [])),
            ("Frameworks & Libraries", skills.get("frameworks", [])),
            ("Tools", skills.get("tools", [])),
            ("Platforms & Cloud", skills.get("platforms", [])),
        ]
        for label, items in skill_lines:
            if not items:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            bold_run = p.add_run(f"{label}: ")
            _set_font(bold_run, 9.5, bold=True)
            val_run = p.add_run(", ".join(items))
            _set_font(val_run, 9.5)

    # Experience
    first_exp = True
    for exp in data.get("experience", []):
        if first_exp:
            _add_section_header(doc, "Experience")
            first_exp = False
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(1)
        title_run = p.add_run(f"{exp.get('title', '')}  ")
        _set_font(title_run, 10, bold=True)
        company_run = p.add_run(f"@ {exp.get('company', '')}")
        _set_font(company_run, 10, color="2C3E7A")

        date_str = f"{exp.get('start_date', '')} – {exp.get('end_date', '')}"
        if exp.get("location"):
            date_str += f"  |  {exp['location']}"
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        date_run = p2.add_run(date_str)
        _set_font(date_run, 9, color="666666")

        for bullet in exp.get("bullets", []):
            _add_bullet(doc, bullet)

    # Projects
    projects = data.get("projects", [])
    if projects:
        _add_section_header(doc, "Projects")
        for proj in projects:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(1)
            name_run = p.add_run(f"{proj.get('name', '')}  ")
            _set_font(name_run, 10, bold=True)
            if proj.get("technologies"):
                tech_run = p.add_run(f"[{', '.join(proj['technologies'][:6])}]")
                _set_font(tech_run, 9, color="555555")
            if proj.get("description"):
                desc_para = doc.add_paragraph()
                desc_para.paragraph_format.left_indent = Inches(0.15)
                desc_para.paragraph_format.space_after = Pt(1)
                _set_font(desc_para.add_run(proj["description"]), 9.5)
            if proj.get("link"):
                link_para = doc.add_paragraph()
                link_para.paragraph_format.left_indent = Inches(0.15)
                _set_font(link_para.add_run(proj["link"]), 9, color="2C3E7A")

    # Education
    first_edu = True
    for edu in data.get("education", []):
        if first_edu:
            _add_section_header(doc, "Education")
            first_edu = False
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        _set_font(p.add_run(f"{edu.get('degree', '')}  "), 10, bold=True)
        _set_font(p.add_run(f"{edu.get('institution', '')}  "), 10, color="2C3E7A")
        _set_font(p.add_run(f"({edu.get('year', '')})"), 9.5, color="666666")
        if edu.get("details"):
            dp = doc.add_paragraph()
            dp.paragraph_format.left_indent = Inches(0.15)
            _set_font(dp.add_run(edu["details"]), 9)

    # Certifications
    certs = data.get("certifications", [])
    if certs:
        _add_section_header(doc, "Certifications")
        for cert in certs:
            _add_bullet(doc, cert)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


# =============================================================================
# Main Composer Class
# =============================================================================

class ResumeComposerV2:
    """Iterative resume composer using LangGraph workflow."""

    def __init__(self, db: DatabaseManager):
        self.db = db
        self.workflow = build_workflow()

    def compose(self, job_id: int) -> str:
        """Compose a tailored resume using the iterative workflow."""
        job = self.db.get_job(job_id)
        if not job:
            raise ValueError(f"Job {job_id} not found")

        gap_analysis_raw = job.get("gap_analysis")
        if not gap_analysis_raw:
            raise ValueError(f"Job {job_id} has no gap analysis. Run 'Assess Resume' first.")

        gap_analysis = (json.loads(gap_analysis_raw) if isinstance(gap_analysis_raw, str)
                        else gap_analysis_raw)

        resume_path = config.MASTER_RESUME_PATH
        if not resume_path.exists():
            raise FileNotFoundError(f"Master resume not found at {resume_path}")

        logger.info("Composing resume (V2 workflow) for job %d: %s @ %s",
                    job_id, job.get("job_title"), job.get("company"))

        # Prepare initial state
        keywords_raw = job.get("keywords") or "[]"
        keywords = (json.loads(keywords_raw) if isinstance(keywords_raw, str) else keywords_raw)

        initial_state: ComposerState = {
            "job_id": job_id,
            "job_title": job.get("job_title", ""),
            "company": job.get("company", ""),
            "job_description": job.get("job_description") or "",
            "keywords": keywords if isinstance(keywords, list) else [],
            "resume_text": _read_docx_text(resume_path)[:8000],
            "gap_analysis": gap_analysis,
            "current_draft": {},
            "assessment": {},
            "revision_count": 0,
            "ats_report": {},
            "final_score": 0,
            "stop_reason": "",
        }

        # Run the workflow
        final_state = self.workflow.invoke(initial_state)

        resume_data = final_state["current_draft"]
        final_score = final_state["final_score"]
        stop_reason = final_state["stop_reason"]
        ats_report = final_state["ats_report"]

        # Generate output path
        company_slug = re.sub(r"[^\w]", "_", job.get("company") or "Company")
        title_slug = re.sub(r"[^\w]", "_", job.get("job_title") or "Role")
        output_path = config.RESUMES_DIR / f"{company_slug}_{title_slug}.docx"

        # Build DOCX
        _build_docx(resume_data, output_path)

        # Log results
        self.db.update_job_gap_analysis(job_id, gap_analysis, final_score)
        self.db.update_job_resume_path(job_id, str(output_path))
        self.db.log_action(job_id, "compose_v2", "completed", {
            "resume_path": str(output_path),
            "final_score": final_score,
            "stop_reason": stop_reason,
            "revision_count": final_state["revision_count"],
            "ats_report": ats_report,
        })

        logger.info("Resume saved to %s | Final score: %d%% | %s",
                    output_path, final_score, stop_reason)

        if ats_report.get("issues"):
            critical = [i for i in ats_report["issues"] if i.get("severity") == "critical"]
            if critical:
                logger.warning("ATS critical issues: %s", critical)

        return str(output_path)

    def get_workflow_diagram(self) -> str:
        """Return ASCII representation of the workflow."""
        return """
        ┌─────────────────┐
        │  generate_draft │
        └────────┬────────┘
                 │
                 ▼
        ┌─────────────────┐
        │     assess      │
        └────────┬────────┘
                 │
        ┌────────▼────────┐
        │   should_revise │─── score >= threshold ───► check_ats
        └────────┬────────┘                             │
                 │                                      ▼
            rounds < max                       ┌─────────────────┐
                 │                             │    finalize     │
                 ▼                             └────────┬────────┘
        ┌─────────────────┐                             │
        │     revise      │                             ▼
        └────────┬────────┘                          (END)
                 │
                 └─────────────────────────────────────┘
                              (loop)
        """
