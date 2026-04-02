"""
Agent 3: Resume Composer
Takes a gap analysis and rewrites the master resume tailored to the specific job.
Outputs an ATS-optimized DOCX file.
"""
import json
import re
from pathlib import Path

import ollama
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import config
from config import setup_logging
from db import DatabaseManager
from agents.assessor import _read_docx_text, _call_llm_with_retry

logger = setup_logging("jobbot.composer")


SYSTEM_PROMPT = """You are an expert resume writer specializing in ATS-optimized resumes for \
software engineers and data/ML/AI engineers. You write compelling, achievement-focused bullet \
points using the STAR method. You never fabricate experience — you reframe and emphasize existing \
experience to highlight relevance."""

USER_PROMPT_TEMPLATE = """Rewrite the following resume to be tailored for the job below. \
Address all critical gaps identified in the analysis without fabricating any experience.

MASTER RESUME:
{resume_text}

TARGET JOB TITLE: {job_title}
TARGET COMPANY: {company}

JOB DESCRIPTION:
{job_description}

GAP ANALYSIS:
{gap_analysis_json}

Rules:
1. Do NOT invent experience, certifications, or skills not implied by the original resume
2. Reorder and reframe existing bullets to highlight relevance to this specific role
3. Incorporate missing keywords naturally into existing bullet points
4. Fix all ATS issues identified in the gap analysis
5. Use strong action verbs: Led, Architected, Optimized, Deployed, Reduced, Increased, Built
6. Quantify achievements where possible
7. Use the recommended_title from the gap analysis as the resume headline

Respond ONLY with a valid JSON object (no markdown, no explanation, no <think> blocks):
{{
  "name": "<full name from master resume>",
  "contact": {{
    "email": "<email>",
    "phone": "<phone>",
    "linkedin": "<linkedin url or username>",
    "github": "<github url or username>",
    "location": "<city, state>"
  }},
  "headline": "<job title / professional headline for this application>",
  "summary": "<2-3 sentence professional summary mentioning the target role>",
  "skills": {{
    "languages": ["<language>"],
    "frameworks": ["<framework>"],
    "tools": ["<tool>"],
    "platforms": ["<platform/cloud>"]
  }},
  "experience": [
    {{
      "title": "<job title>",
      "company": "<company name>",
      "location": "<location>",
      "start_date": "<Mon YYYY>",
      "end_date": "<Mon YYYY or Present>",
      "bullets": ["<rewritten bullet>"]
    }}
  ],
  "education": [
    {{
      "degree": "<degree>",
      "institution": "<school name>",
      "year": "<graduation year>",
      "details": "<GPA or relevant coursework if notable>"
    }}
  ],
  "projects": [
    {{
      "name": "<project name>",
      "description": "<1-2 sentence description>",
      "technologies": ["<tech>"],
      "link": "<url or empty string>"
    }}
  ],
  "certifications": ["<certification name and year>"]
}}"""


# ── DOCX builder ──────────────────────────────────────────────────────────────

def _add_horizontal_rule(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2C3E7A")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_font(run, size: float, bold: bool = False, color: str | None = None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)


def _add_section_header(doc: Document, text: str):
    para = doc.add_paragraph()
    run = para.add_run(text.upper())
    _set_font(run, 10.5, bold=True, color="2C3E7A")
    _add_horizontal_rule(para)
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(2)
    return para


def _add_bullet(doc: Document, text: str):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.15)
    para.paragraph_format.space_after = Pt(1)
    run = para.add_run(f"\u2022  {text}")
    _set_font(run, 9.5)
    return para


def _build_docx(data: dict, output_path: Path) -> Path:
    doc = Document()

    # Ensure all top-level keys have safe defaults so missing LLM keys
    # produce a graceful skip rather than a KeyError crash
    contact = data.get("contact") or {}
    summary = data.get("summary") or ""
    skills = data.get("skills") or {}
    experience = data.get("experience") or []
    education = data.get("education") or []
    certifications = data.get("certifications") or []
    projects = data.get("projects") or []

    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    doc.styles["Normal"].paragraph_format.space_before = Pt(0)
    doc.styles["Normal"].paragraph_format.space_after = Pt(0)

    # Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(data.get("name", ""))
    _set_font(name_run, 16, bold=True, color="1A1A2E")
    name_para.paragraph_format.space_after = Pt(2)

    # Headline
    if data.get("headline"):
        hl_para = doc.add_paragraph()
        hl_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hl_run = hl_para.add_run(data["headline"])
        _set_font(hl_run, 10.5, color="2C3E7A")
        hl_para.paragraph_format.space_after = Pt(3)

    # Contact line
    contact_parts = []
    for key in ("email", "phone", "location"):
        if contact.get(key):
            contact_parts.append(contact[key])
    if contact.get("linkedin"):
        contact_parts.append(f"LinkedIn: {contact['linkedin']}")
    if contact.get("github"):
        contact_parts.append(f"GitHub: {contact['github']}")
    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run("  |  ".join(contact_parts))
        _set_font(contact_run, 9, color="444444")
        contact_para.paragraph_format.space_after = Pt(4)

    # Summary
    if summary:
        _add_section_header(doc, "Professional Summary")
        p = doc.add_paragraph()
        r = p.add_run(summary)
        _set_font(r, 9.5)
        p.paragraph_format.space_after = Pt(3)

    # Skills
    if skills and any(skills.values()):
        _add_section_header(doc, "Technical Skills")
        skill_lines = [
            ("Languages", skills.get("languages") or []),
            ("Frameworks & Libraries", skills.get("frameworks") or []),
            ("Tools", skills.get("tools") or []),
            ("Platforms & Cloud", skills.get("platforms") or []),
        ]
        for label, items in skill_lines:
            if not items:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            bold_run = p.add_run(f"{label}: ")
            _set_font(bold_run, 9.5, bold=True)
            val_run = p.add_run(", ".join(items))
            _set_font(val_run, 9.5)

    # Experience
    for idx, exp in enumerate(experience):
        if idx == 0:
            _add_section_header(doc, "Experience")
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(1)
        title_run = p.add_run(f"{exp.get('title', '')}  ")
        _set_font(title_run, 10, bold=True)
        company_run = p.add_run(f"@ {exp.get('company', '')}")
        _set_font(company_run, 10, color="2C3E7A")

        date_str = f"{exp.get('start_date', '')} – {exp.get('end_date', '')}"
        if exp.get("location"):
            date_str += f"  |  {exp['location']}"
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        date_run = p2.add_run(date_str)
        _set_font(date_run, 9, color="666666")

        for bullet in exp.get("bullets", []):
            _add_bullet(doc, bullet)

    # Projects
    if projects:
        _add_section_header(doc, "Projects")
        for proj in projects:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(1)
            name_run = p.add_run(f"{proj.get('name', '')}  ")
            _set_font(name_run, 10, bold=True)
            if proj.get("technologies"):
                tech_run = p.add_run(f"[{', '.join(proj['technologies'][:6])}]")
                _set_font(tech_run, 9, color="555555")
            if proj.get("description"):
                desc_para = doc.add_paragraph()
                desc_para.paragraph_format.left_indent = Inches(0.15)
                desc_para.paragraph_format.space_after = Pt(1)
                _set_font(desc_para.add_run(proj["description"]), 9.5)
            if proj.get("link"):
                link_para = doc.add_paragraph()
                link_para.paragraph_format.left_indent = Inches(0.15)
                _set_font(link_para.add_run(proj["link"]), 9, color="2C3E7A")

    # Education
    for idx, edu in enumerate(education):
        if idx == 0:
            _add_section_header(doc, "Education")
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        _set_font(p.add_run(f"{edu.get('degree', '')}  "), 10, bold=True)
        _set_font(p.add_run(f"{edu.get('institution', '')}  "), 10, color="2C3E7A")
        _set_font(p.add_run(f"({edu.get('year', '')})"), 9.5, color="666666")
        if edu.get("details"):
            dp = doc.add_paragraph()
            dp.paragraph_format.left_indent = Inches(0.15)
            _set_font(dp.add_run(edu["details"]), 9)

    # Certifications
    if certifications:
        _add_section_header(doc, "Certifications")
        for cert in certifications:
            _add_bullet(doc, cert)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


# ── Main class ────────────────────────────────────────────────────────────────

class ResumeComposer:
    def __init__(self, db: DatabaseManager):
        self.db = db

    def compose(self, job_id: int) -> str:
        """Compose a tailored resume for the given job. Returns path to DOCX."""
        job = self.db.get_job(job_id)
        if not job:
            raise ValueError(f"Job {job_id} not found")

        gap_analysis_raw = job.get("gap_analysis")
        if not gap_analysis_raw:
            raise ValueError(f"Job {job_id} has no gap analysis. Run 'Assess Resume' first.")

        gap_analysis = json.loads(gap_analysis_raw) if isinstance(gap_analysis_raw, str) else gap_analysis_raw

        resume_path = config.MASTER_RESUME_PATH
        if not resume_path.exists():
            raise FileNotFoundError(f"Master resume not found at {resume_path}")

        logger.info("Composing resume for job %d: %s @ %s", job_id, job.get("job_title"), job.get("company"))
        resume_text = _read_docx_text(resume_path)[:8000]

        user_prompt = USER_PROMPT_TEMPLATE.format(
            resume_text=resume_text,
            job_title=job.get("job_title", ""),
            company=job.get("company", ""),
            job_description=(job.get("job_description") or "")[:5000],
            gap_analysis_json=json.dumps(gap_analysis, indent=2),
        )

        resume_data, usage = _call_llm_with_retry(SYSTEM_PROMPT, user_prompt, max_tokens=4000)
        logger.debug("Compose LLM usage — input: %d tokens, output: %d tokens",
                     usage.get("input_tokens", 0), usage.get("output_tokens", 0))

        company_slug = re.sub(r"[^\w]", "_", job.get("company") or "Company")
        title_slug = re.sub(r"[^\w]", "_", job.get("job_title") or "Role")
        output_path = config.RESUMES_DIR / f"{company_slug}_{title_slug}.docx"

        _build_docx(resume_data, output_path)

        self.db.update_job_resume_path(job_id, str(output_path))
        self.db.save_resume_version(job_id, str(output_path), usage)
        self.db.log_action(job_id, "compose", "completed",
                           {"resume_path": str(output_path), "tokens": usage})

        logger.info("Resume saved to %s", output_path)
        return str(output_path)
